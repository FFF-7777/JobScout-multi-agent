"""文档解析：从 PDF / DOCX / MD / TXT 提取文本；从 Excel / CSV 提取岗位列表。"""
from __future__ import annotations

import io

import pandas as pd


def parse_resume_bytes(filename: str, data: bytes) -> str:
    """根据文件后缀解析简历文本。"""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _parse_pdf(data)
    if name.endswith(".docx"):
        return _parse_docx(data)
    if name.endswith((".md", ".markdown", ".txt")):
        return data.decode("utf-8", errors="ignore")
    # 兜底按文本处理
    return data.decode("utf-8", errors="ignore")


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(parts).strip()


def _parse_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    # 也读取表格内容
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


# 岗位表格字段的中英文别名映射
_JOB_FIELD_ALIASES = {
    "company_name": ["company_name", "company", "公司", "公司名称"],
    "job_title": ["job_title", "title", "岗位", "职位", "岗位名称", "职位名称"],
    "city": ["city", "城市", "工作城市", "地点"],
    "salary": ["salary", "薪资", "薪水", "待遇"],
    "jd_text": ["jd_text", "jd", "岗位描述", "职位描述", "岗位职责", "描述"],
    "job_url": ["job_url", "url", "链接", "岗位链接"],
    "source": ["source", "来源", "平台"],
    "education": ["education", "学历"],
    "experience": ["experience", "经验", "工作经验"],
}


def parse_jobs_table(filename: str, data: bytes) -> list[dict]:
    """解析 Excel / CSV 岗位列表，返回规范化的岗位字典列表。"""
    name = filename.lower()
    if name.endswith(".csv"):
        df = pd.read_csv(io.BytesIO(data))
    elif name.endswith((".xlsx", ".xls")):
        df = pd.read_excel(io.BytesIO(data))
    else:
        raise ValueError("仅支持 .csv / .xlsx / .xls 岗位表格")

    # 建立列名 -> 标准字段 的映射
    col_map: dict[str, str] = {}
    for col in df.columns:
        key = str(col).strip().lower()
        for std, aliases in _JOB_FIELD_ALIASES.items():
            if key in [a.lower() for a in aliases]:
                col_map[col] = std
                break

    jobs: list[dict] = []
    for _, row in df.iterrows():
        job: dict = {
            "source": "excel",
            "company_name": "",
            "job_title": "",
            "city": "",
            "salary": "",
            "education": "",
            "experience": "",
            "jd_text": "",
            "job_url": "",
        }
        for col, std in col_map.items():
            val = row[col]
            if pd.isna(val):
                continue
            job[std] = str(val).strip()
        # 若没有 jd_text，用其它字段拼一个，便于后续 Job Agent 解析
        if not job["jd_text"]:
            job["jd_text"] = "；".join(
                f"{k}:{v}" for k, v in job.items() if v and k != "jd_text"
            )
        # 跳过完全空行
        if any(job[k] for k in ("company_name", "job_title", "jd_text")):
            jobs.append(job)
    return jobs
